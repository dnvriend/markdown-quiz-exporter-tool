"""DOCX exporter for quiz questions.

This module exports parsed quiz questions to Microsoft Word DOCX format.
Supports markdown formatting including bold, italic, and code blocks.

Note: This code was generated with assistance from AI coding tools
and has been reviewed and tested by a human.
"""

import re
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt

from markdown_quiz_exporter_tool.quiz_parser import Answer, Question

# Unicode checkbox characters
CHECKBOX_CHECKED = "\u2611"  # ☑
CHECKBOX_UNCHECKED = "\u2610"  # ☐

# Font settings
FONT_NAME = "Times New Roman"
FONT_SIZE = Pt(12)
CODE_FONT_NAME = "Courier New"


class DocxExporter:
    """Exporter for Word DOCX format."""

    def __init__(self, questions: list[Question]) -> None:
        """Initialize the exporter with parsed questions.

        Args:
            questions: List of Question objects to export
        """
        self.questions = questions
        self.document = Document()
        self._setup_styles()

    def _setup_styles(self) -> None:
        """Setup document styles for consistent formatting."""
        # Set default font for Normal style
        style = self.document.styles["Normal"]
        font = style.font
        font.name = FONT_NAME
        font.size = FONT_SIZE

        # Create Code style for code blocks
        if "Code" not in self.document.styles:
            code_style = self.document.styles.add_style("Code", WD_STYLE_TYPE.PARAGRAPH)
            code_font = code_style.font
            code_font.name = CODE_FONT_NAME
            code_font.size = Pt(10)
            code_style.paragraph_format.left_indent = Pt(20)
            code_style.paragraph_format.space_before = Pt(6)
            code_style.paragraph_format.space_after = Pt(6)

    def export(self, output_path: Path) -> int:
        """Export questions to DOCX format.

        Args:
            output_path: Path where the DOCX file will be written

        Returns:
            Number of questions exported

        Raises:
            IOError: If writing to file fails
        """
        for i, question in enumerate(self.questions):
            self._add_question(question, i + 1)

            # Add horizontal line separator (except after last question)
            if i < len(self.questions) - 1:
                self._add_horizontal_line()

        self.document.save(str(output_path))
        return len(self.questions)

    def _add_question(self, question: Question, number: int) -> None:
        """Add a question to the document.

        Args:
            question: Question object to add
            number: Question number (1-indexed)
        """
        # Add question text with number
        question_para = self.document.add_paragraph()
        run = question_para.add_run(f"{number}. ")
        run.bold = True
        run.font.name = FONT_NAME
        run.font.size = FONT_SIZE

        # Add question text with markdown formatting
        self._add_formatted_text(question_para, question.text)

        # Add answers
        for answer in question.answers:
            self._add_answer(answer)

        # Add reason section if present
        if question.reason.strip():
            self._add_reason(question.reason)

        # Add spacing after question
        self.document.add_paragraph()

    def _add_answer(self, answer: Answer) -> None:
        """Add an answer option to the document.

        Args:
            answer: Answer object to add
        """
        para = self.document.add_paragraph()
        para.paragraph_format.left_indent = Pt(20)

        # Add checkbox
        checkbox = CHECKBOX_CHECKED if answer.is_correct else CHECKBOX_UNCHECKED
        run = para.add_run(f"    {checkbox} ")
        run.font.name = FONT_NAME
        run.font.size = FONT_SIZE

        # Add answer text with markdown formatting
        self._add_formatted_text(para, answer.text)

    def _add_reason(self, reason: str) -> None:
        """Add the reason/explanation section.

        Args:
            reason: Explanation text
        """
        # Add "Reason:" label
        para = self.document.add_paragraph()
        para.paragraph_format.space_before = Pt(12)

        label_run = para.add_run("Reason: ")
        label_run.bold = True
        label_run.italic = True
        label_run.font.name = FONT_NAME
        label_run.font.size = FONT_SIZE

        # Add reason text in italic with markdown formatting
        self._add_formatted_text(para, reason, base_italic=True)

    def _add_formatted_text(
        self, paragraph: Any, text: str, base_italic: bool = False
    ) -> None:
        """Add text with markdown formatting to a paragraph.

        Handles bold (**text**), italic (*text*), inline code (`code`),
        and code blocks (```code```).

        Args:
            paragraph: Paragraph to add text to
            text: Markdown text to format
            base_italic: If True, apply italic as base style
        """
        # Check for code blocks first
        code_block_pattern = re.compile(r"```(\w*)\n?(.*?)```", re.DOTALL)
        parts = code_block_pattern.split(text)

        if len(parts) > 1:
            # Has code blocks - process each part
            i = 0
            while i < len(parts):
                if i % 3 == 0:
                    # Regular text
                    if parts[i].strip():
                        self._add_inline_formatted_text(paragraph, parts[i], base_italic)
                elif i % 3 == 2:
                    # Code block content
                    self._add_code_block(parts[i])
                i += 1
        else:
            # No code blocks - just inline formatting
            self._add_inline_formatted_text(paragraph, text, base_italic)

    def _add_inline_formatted_text(
        self, paragraph: Any, text: str, base_italic: bool = False
    ) -> None:
        """Add text with inline markdown formatting.

        Args:
            paragraph: Paragraph to add text to
            text: Text with inline markdown
            base_italic: If True, apply italic as base style
        """
        # Pattern to match bold, italic, and inline code
        # Order matters: bold (**) before italic (*) to avoid conflicts
        pattern = re.compile(
            r"(\*\*(.+?)\*\*)"  # Bold: **text**
            r"|(\*(.+?)\*)"  # Italic: *text*
            r"|(`([^`]+)`)"  # Inline code: `code`
        )

        last_end = 0
        for match in pattern.finditer(text):
            # Add text before match
            if match.start() > last_end:
                run = paragraph.add_run(text[last_end : match.start()])
                run.font.name = FONT_NAME
                run.font.size = FONT_SIZE
                if base_italic:
                    run.italic = True

            # Determine which pattern matched
            if match.group(2):  # Bold
                run = paragraph.add_run(match.group(2))
                run.bold = True
                run.font.name = FONT_NAME
                run.font.size = FONT_SIZE
                if base_italic:
                    run.italic = True
            elif match.group(4):  # Italic
                run = paragraph.add_run(match.group(4))
                run.italic = True
                run.font.name = FONT_NAME
                run.font.size = FONT_SIZE
            elif match.group(6):  # Inline code
                run = paragraph.add_run(match.group(6))
                run.font.name = CODE_FONT_NAME
                run.font.size = FONT_SIZE
                # Add light gray background for inline code
                self._set_run_highlight(run, "lightGray")

            last_end = match.end()

        # Add remaining text
        if last_end < len(text):
            run = paragraph.add_run(text[last_end:])
            run.font.name = FONT_NAME
            run.font.size = FONT_SIZE
            if base_italic:
                run.italic = True

    def _add_code_block(self, code: str) -> None:
        """Add a code block with monospace font and shading.

        Args:
            code: Code content
        """
        code = code.strip()
        if not code:
            return

        para = self.document.add_paragraph()
        para.style = self.document.styles["Code"]

        # Add shading to the paragraph
        self._set_paragraph_shading(para, "E8E8E8")  # Light gray

        for line in code.split("\n"):
            if line:
                run = para.add_run(line)
                run.font.name = CODE_FONT_NAME
                run.font.size = Pt(10)
            para.add_run("\n")

    def _set_paragraph_shading(self, paragraph: Any, color: str) -> None:
        """Set background shading for a paragraph.

        Args:
            paragraph: Paragraph to shade
            color: Hex color code (without #)
        """
        shading_elm = OxmlElement("w:shd")
        shading_elm.set(qn("w:fill"), color)
        paragraph._p.get_or_add_pPr().append(shading_elm)

    def _set_run_highlight(self, run: Any, color: str) -> None:
        """Set highlight color for a run.

        Args:
            run: Run to highlight
            color: Word highlight color name (e.g., 'lightGray', 'yellow')
        """
        highlight = OxmlElement("w:highlight")
        highlight.set(qn("w:val"), color)
        run._r.get_or_add_rPr().append(highlight)

    def _add_horizontal_line(self) -> None:
        """Add a horizontal line separator."""
        para = self.document.add_paragraph()
        para.paragraph_format.space_before = Pt(12)
        para.paragraph_format.space_after = Pt(12)

        # Add bottom border to paragraph
        p_pr = para._p.get_or_add_pPr()
        p_bdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "6")
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), "auto")
        p_bdr.append(bottom)
        p_pr.append(p_bdr)


def export_to_docx(questions: list[Question], output_path: Path) -> int:
    """Export questions to Word DOCX format.

    Args:
        questions: List of Question objects to export
        output_path: Path where the DOCX file will be written

    Returns:
        Number of questions exported

    Raises:
        IOError: If writing to file fails
    """
    exporter = DocxExporter(questions)
    return exporter.export(output_path)
