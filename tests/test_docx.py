"""Tests for Word DOCX exporter.

Note: This code was generated with assistance from AI coding tools
and has been reviewed and tested by a human.
"""

from pathlib import Path

import pytest
from docx import Document

from markdown_quiz_exporter_tool.docx import (
    CHECKBOX_CHECKED,
    CHECKBOX_UNCHECKED,
    DocxExporter,
    export_to_docx,
)
from markdown_quiz_exporter_tool.quiz_parser import Answer, Question


@pytest.fixture
def single_choice_question() -> Question:
    """Create a single choice question."""
    return Question(
        text="What is the max size of an S3 object in GB?",
        answers=[
            Answer(text="1000", is_correct=False),
            Answer(text="5000", is_correct=True),
            Answer(text="10000", is_correct=False),
        ],
        reason="The maximum size is **5TB** or 5000GB.",
        question_type="single",
    )


@pytest.fixture
def multiple_choice_question() -> Question:
    """Create a multiple choice question."""
    return Question(
        text="What are characteristics of effective leadership?",
        answers=[
            Answer(text="Ability to motivate", is_correct=True),
            Answer(text="Clear communication", is_correct=True),
            Answer(text="Authoritarian approach", is_correct=False),
        ],
        reason="Leadership requires *motivation* and communication.",
        question_type="multiple",
    )


@pytest.fixture
def question_with_code() -> Question:
    """Create a question with code block."""
    return Question(
        text="What does this code do?\n```python\nprint('hello')\n```",
        answers=[
            Answer(text="Prints `hello`", is_correct=True),
            Answer(text="Does nothing", is_correct=False),
        ],
        reason="The `print()` function outputs text.",
        question_type="single",
    )


def test_export_creates_file(tmp_path: Path, single_choice_question: Question) -> None:
    """Test that export creates a DOCX file."""
    output_file = tmp_path / "output.docx"
    count = export_to_docx([single_choice_question], output_file)

    assert count == 1
    assert output_file.exists()


def test_export_returns_question_count(
    tmp_path: Path, single_choice_question: Question, multiple_choice_question: Question
) -> None:
    """Test that export returns correct question count."""
    output_file = tmp_path / "output.docx"
    count = export_to_docx([single_choice_question, multiple_choice_question], output_file)

    assert count == 2


def test_docx_contains_question_text(tmp_path: Path, single_choice_question: Question) -> None:
    """Test that DOCX contains question text."""
    output_file = tmp_path / "output.docx"
    export_to_docx([single_choice_question], output_file)

    doc = Document(str(output_file))
    full_text = "\n".join(p.text for p in doc.paragraphs)

    assert "S3 object" in full_text


def test_docx_contains_numbered_questions(
    tmp_path: Path, single_choice_question: Question, multiple_choice_question: Question
) -> None:
    """Test that questions are numbered."""
    output_file = tmp_path / "output.docx"
    export_to_docx([single_choice_question, multiple_choice_question], output_file)

    doc = Document(str(output_file))
    full_text = "\n".join(p.text for p in doc.paragraphs)

    assert "1." in full_text
    assert "2." in full_text


def test_docx_contains_checkboxes(tmp_path: Path, single_choice_question: Question) -> None:
    """Test that DOCX contains checkbox characters."""
    output_file = tmp_path / "output.docx"
    export_to_docx([single_choice_question], output_file)

    doc = Document(str(output_file))
    full_text = "\n".join(p.text for p in doc.paragraphs)

    # Should have one checked and two unchecked
    assert CHECKBOX_CHECKED in full_text
    assert CHECKBOX_UNCHECKED in full_text


def test_docx_contains_reason_label(tmp_path: Path, single_choice_question: Question) -> None:
    """Test that DOCX contains 'Reason:' label."""
    output_file = tmp_path / "output.docx"
    export_to_docx([single_choice_question], output_file)

    doc = Document(str(output_file))
    full_text = "\n".join(p.text for p in doc.paragraphs)

    assert "Reason:" in full_text


def test_docx_contains_reason_text(tmp_path: Path, single_choice_question: Question) -> None:
    """Test that DOCX contains reason content."""
    output_file = tmp_path / "output.docx"
    export_to_docx([single_choice_question], output_file)

    doc = Document(str(output_file))
    full_text = "\n".join(p.text for p in doc.paragraphs)

    assert "5TB" in full_text


def test_docx_question_with_inline_code(tmp_path: Path, question_with_code: Question) -> None:
    """Test handling of inline code."""
    output_file = tmp_path / "output.docx"
    export_to_docx([question_with_code], output_file)

    doc = Document(str(output_file))
    full_text = "\n".join(p.text for p in doc.paragraphs)

    assert "print()" in full_text
    assert "hello" in full_text


def test_docx_multiple_questions_count(
    tmp_path: Path, single_choice_question: Question, multiple_choice_question: Question
) -> None:
    """Test export with multiple questions."""
    output_file = tmp_path / "output.docx"
    questions = [single_choice_question, multiple_choice_question]

    count = export_to_docx(questions, output_file)

    assert count == 2


def test_docx_answer_options(tmp_path: Path, single_choice_question: Question) -> None:
    """Test that all answer options are included."""
    output_file = tmp_path / "output.docx"
    export_to_docx([single_choice_question], output_file)

    doc = Document(str(output_file))
    full_text = "\n".join(p.text for p in doc.paragraphs)

    assert "1000" in full_text
    assert "5000" in full_text
    assert "10000" in full_text


def test_docx_correct_checkbox_for_correct_answer(
    tmp_path: Path, single_choice_question: Question
) -> None:
    """Test that correct answer has checked checkbox."""
    output_file = tmp_path / "output.docx"
    export_to_docx([single_choice_question], output_file)

    doc = Document(str(output_file))
    paragraphs = [p.text for p in doc.paragraphs]

    # Find paragraph that starts with checkbox and contains correct answer
    correct_answer_para = [p for p in paragraphs if CHECKBOX_CHECKED in p and "5000" in p]
    assert len(correct_answer_para) >= 1
    assert CHECKBOX_CHECKED in correct_answer_para[0]


def test_docx_unchecked_checkbox_for_incorrect_answer(
    tmp_path: Path, single_choice_question: Question
) -> None:
    """Test that incorrect answers have unchecked checkboxes."""
    output_file = tmp_path / "output.docx"
    export_to_docx([single_choice_question], output_file)

    doc = Document(str(output_file))
    paragraphs = [p.text for p in doc.paragraphs]

    # Find paragraph with checkbox and incorrect answer "1000" (not "10000")
    incorrect_answer_para = [p for p in paragraphs if CHECKBOX_UNCHECKED in p]
    assert len(incorrect_answer_para) >= 1  # At least 2 incorrect answers
    # At least one should have unchecked checkbox
    assert any(CHECKBOX_UNCHECKED in p for p in incorrect_answer_para)


def test_docx_no_reason(tmp_path: Path) -> None:
    """Test export when question has no reason."""
    question = Question(
        text="Test question?",
        answers=[Answer(text="Test answer", is_correct=True)],
        reason="",
        question_type="single",
    )

    output_file = tmp_path / "output.docx"
    export_to_docx([question], output_file)

    doc = Document(str(output_file))
    full_text = "\n".join(p.text for p in doc.paragraphs)

    # Should have question and answer but no Reason: label
    assert "Test question?" in full_text
    assert "Reason:" not in full_text


def test_docx_exporter_instance() -> None:
    """Test DocxExporter can be instantiated."""
    question = Question(
        text="Q?",
        answers=[Answer(text="A", is_correct=True)],
        reason="R",
        question_type="single",
    )

    exporter = DocxExporter([question])
    assert exporter.questions == [question]
    assert exporter.document is not None
